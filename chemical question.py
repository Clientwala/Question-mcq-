#!/usr/bin/env python3
"""
Chemistry/Engineering - SAIL Questions
Q76-100 with Detailed Solutions
FINAL CENTERED FORMAT (No a,b,c,d prefix in options)
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_element = OxmlElement(f'w:{edge}')
        edge_element.set(qn('w:val'), 'single')
        edge_element.set(qn('w:sz'), '4')
        edge_element.set(qn('w:space'), '0')
        edge_element.set(qn('w:color'), '000000')
        tcBorders.append(edge_element)
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=120, bottom=120, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_column_width(column, width_cm):
    for cell in column.cells:
        cell.width = Cm(width_cm)

def add_text_to_cell(cell, text_parts):
    cell.text = ''
    if isinstance(text_parts, str):
        text_parts = [text_parts]
    for idx, part in enumerate(text_parts):
        if idx > 0:
            paragraph = cell.add_paragraph()
        else:
            paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if idx > 0:
            paragraph.paragraph_format.space_before = Pt(3)
        run = paragraph.add_run(part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True

def add_question_table(doc, question_parts, options, correct_option_idx, solution_parts, has_diagram=False):
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    set_column_width(table.columns[0], 1.5)
    set_column_width(table.columns[1], 8.5)
    set_column_width(table.columns[2], 3.0)
    
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell, top=120, bottom=120, left=100, right=100)
    
    table.rows[0].cells[0].text = "Question"
    cell = table.rows[0].cells[1].merge(table.rows[0].cells[2])
    if has_diagram:
        add_text_to_cell(cell, question_parts + ['', '[DIAGRAM PRESENT - See PDF]'])
    else:
        add_text_to_cell(cell, question_parts)
    
    table.rows[1].cells[0].text = "Type"
    table.rows[1].cells[1].text = "Multiple_choice"
    table.rows[1].cells[1].merge(table.rows[1].cells[2])
    
    for i in range(4):
        row_idx = i + 2
        table.rows[row_idx].cells[0].text = "Option"
        table.rows[row_idx].cells[1].text = options[i] if i < len(options) else ""
        if i == correct_option_idx:
            table.rows[row_idx].cells[2].text = "correct"
        else:
            table.rows[row_idx].cells[2].text = "incorrect"
    
    table.rows[6].cells[0].text = "Solution"
    cell = table.rows[6].cells[1].merge(table.rows[6].cells[2])
    add_text_to_cell(cell, solution_parts)
    
    table.rows[7].cells[0].text = "Marks"
    table.rows[7].cells[1].text = "1"
    table.rows[7].cells[2].text = "0.25"
    
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                    run.font.bold = True

questions_data = [
    # Q76 - Safety Instrumented System
    {'question': ['Safety Instrumented Systems (SIS) are primarily designed to:'], 
     'options': ['Maximize production rate', 'Reduce process downtime', 'Bring the process to a safe state during abnormal conditions', 'Optimize energy consumption'], 
     'correct_idx': 2, 
     'solution': ['SIS is an independent protection layer designed to prevent or mitigate hazardous events.', 'Purpose: Automatically bring process to safe state when dangerous conditions detected', 'Components: Sensors, logic solver, final control elements (shutdown valves, emergency stops)', 'Different from basic process control system (BPCS) which optimizes normal operation', 'Safety Integrity Level (SIL): Measures reliability (SIL 1 to SIL 4, higher = more reliable)', 'Example: High pressure → SIS shuts down compressor and opens relief valve. Critical in chemical plants, oil refineries, nuclear facilities for protecting people, environment, and assets.'], 
     'diagram': False},
    
    # Q77 - Piping and Instrumentation Diagram
    {'question': ['In a P&ID (Piping and Instrumentation Diagram), a control valve is typically represented by:'], 
     'options': ['A circle', 'A square with diagonal line', 'A triangle', 'A valve symbol with actuator'], 
     'correct_idx': 3, 
     'solution': ['P&ID is a detailed schematic showing piping, equipment, and instrumentation in a process plant.', 'Control valve symbol: Valve body (typically triangular) with actuator symbol on top', 'Actuator types: Pneumatic (shown with air signal line), Electric (motor symbol), Hydraulic', 'Fail position indicated: FC (fail closed), FO (fail open), FL (fail last position)', 'Tag example: FCV-101 (Flow Control Valve, equipment number 101)', 'P&IDs also show: Process equipment, piping sizes, control loops, interlocks, safety devices', 'Standard: ISA (International Society of Automation) symbols. Essential document for plant operation, maintenance, and safety analysis.'], 
     'diagram': False},
    
    # Q78 - Material of Construction
    {'question': ['Stainless steel (SS316) is preferred over carbon steel for handling:'], 
     'options': ['High temperature steam', 'Corrosive chemicals like acids', 'Inert gases', 'Pure water only'], 
     'correct_idx': 1, 
     'solution': ['Material selection depends on service conditions: corrosion, temperature, pressure, cost.', 'SS316 (18% Cr, 12% Ni, 2-3% Mo): Excellent corrosion resistance, especially to chlorides and acids', 'Carbon steel: Economical but corrodes easily with acids, chlorides, oxidizing conditions', 'SS316 applications: Pharmaceutical equipment, food processing, marine applications, chemical handling', 'Molybdenum in SS316 improves pitting resistance in chloride environments', 'Other materials: Hastelloy (severe corrosion), Titanium (high strength, corrosion), FRP (Fiber Reinforced Plastic)', 'Trade-off: SS316 costs 3-5x more than carbon steel. Selection based on life cycle cost and safety requirements.'], 
     'diagram': False},
    
    # Q79 - Pump Cavitation
    {'question': ['Cavitation in a pump occurs when:'], 
     'options': ['Discharge pressure is too high', 'Suction pressure drops below vapor pressure of liquid', 'Flow rate exceeds design', 'Temperature is too low'], 
     'correct_idx': 1, 
     'solution': ['Cavitation: Formation and collapse of vapor bubbles in liquid due to low pressure.', 'Occurs when local pressure falls below vapor pressure (Plocal < Pvapor)', 'Typically at pump suction or impeller inlet (lowest pressure point)', 'Bubble collapse near metal surface causes: Noise, vibration, erosion damage, reduced performance', 'NPSH (Net Positive Suction Head): Margin above vapor pressure. NPSH_available > NPSH_required prevents cavitation.', 'Prevention: (1) Increase suction pressure, (2) Reduce suction line losses, (3) Lower pump elevation, (4) Use larger suction pipe, (5) Reduce liquid temperature', 'Symptoms: Rattling sound, vibration, reduced head. Cavitation damage can destroy impeller in weeks.'], 
     'diagram': False},
    
    # Q80 - Compressor Types
    {'question': ['For high pressure ratio (> 10:1) gas compression, which type is most suitable?'], 
     'options': ['Single-stage centrifugal', 'Multi-stage centrifugal or reciprocating', 'Axial compressor', 'Liquid ring compressor'], 
     'correct_idx': 1, 
     'solution': ['Compression types: Positive displacement (reciprocating, rotary) and Dynamic (centrifugal, axial)', 'Single-stage centrifugal: Pressure ratio typically 1.5-4:1. Higher ratios cause excessive temperature rise and reduced efficiency.', 'Multi-stage centrifugal: 3-8 stages with intercooling. Can achieve 15:1 overall. Used for large flows.', 'Reciprocating: Can achieve very high ratios (> 100:1) with staging and intercooling. Lower flow rates.', 'For pressure ratio > 10:1: Use multi-stage centrifugal OR multi-stage reciprocating with intercoolers', 'Intercooling between stages: Reduces power consumption, prevents overheating, approaches isothermal compression. Applications: Natural gas processing, refrigeration, air separation.'], 
     'diagram': False},
    
    # Q81 - Distillation Column Flooding
    {'question': ['Flooding in a distillation column occurs due to:'], 
     'options': ['Too low vapor velocity', 'Excessive liquid holdup preventing vapor flow', 'Insufficient reflux', 'Low tray efficiency'], 
     'correct_idx': 1, 
     'solution': ['Flooding: Condition where excessive liquid accumulates on trays, preventing normal vapor-liquid contact.', 'Causes: (1) Vapor velocity too high (entrainment flooding - liquid carried upward), (2) Liquid rate too high (downcomer flooding - liquid backup)', 'At flooding: Pressure drop increases sharply, liquid level rises, separation efficiency drops', 'Vapor velocity at flooding ≈ 80-85% of velocity calculated by Fair correlation', 'Design: Operate at 70-80% of flooding velocity for safety margin and turndown capability', 'Symptoms: Erratic operation, high differential pressure, liquid carryover to overhead', 'Prevention: Proper tray design (hole area, downcomer area), anti-foam agents. Flooding limits column capacity and throughput.'], 
     'diagram': False},
    
    # Q82 - Packed Column vs Tray Column
    {'question': ['Compared to tray columns, packed columns have:'], 
     'options': ['Higher pressure drop per theoretical stage', 'Lower pressure drop per theoretical stage', 'Higher liquid holdup', 'Cannot handle fouling services'], 
     'correct_idx': 1, 
     'solution': ['Packed columns: Random or structured packing instead of trays. Continuous vapor-liquid contact.', 'Advantages: (1) Lower pressure drop (30-50% less than trays), (2) Lower liquid holdup, (3) Better for vacuum/corrosive services, (4) Lighter weight', 'Disadvantages: (1) Difficult to clean (no access), (2) Channeling and maldistribution possible, (3) Foaming problems', 'Tray columns: Discrete stages with better mechanical design access', 'Selection criteria: Packing for vacuum, low ΔP, corrosive. Trays for fouling, large diameter, easier maintenance.', 'Modern random packings: Pall rings, Raschig rings, Berl saddles. Structured packing: Higher efficiency but more expensive. Used in distillation, absorption, stripping.'], 
     'diagram': False},
    
    # Q83 - Heat Exchanger Fouling
    {'question': ['Fouling in heat exchangers reduces:'], 
     'options': ['Pressure drop', 'Heat transfer coefficient', 'Flow velocity', 'Tube diameter'], 
     'correct_idx': 1, 
     'solution': ['Fouling: Deposition of unwanted material on heat transfer surfaces (scale, biological growth, corrosion products, particulates)', 'Effect: Adds thermal resistance, reducing overall heat transfer coefficient U', '1/U = 1/hi + Rf,i + Rwall + Rf,o + 1/ho where Rf = fouling resistance', 'Consequences: (1) Reduced heat transfer (lower duty), (2) Increased pressure drop, (3) More frequent cleaning, (4) Reduced run length', 'Types: Crystallization (hardness salts), particulate (suspended solids), biological (algae, bacteria), corrosion, chemical reaction', 'Mitigation: (1) Proper velocity (2-3 m/s minimum), (2) Water treatment, (3) Regular cleaning, (4) Fouling allowance in design (20-30%)', 'Fouling factor typically 0.0002-0.001 m²K/W depending on service. Major cause of heat exchanger degradation.'], 
     'diagram': False},
    
    # Q84 - Shell and Tube Designations
    {'question': ['A 1-2 heat exchanger designation means:'], 
     'options': ['1 shell pass, 2 tube passes', '2 shell passes, 1 tube pass', '1 tube, 2 shells', '1 inlet, 2 outlets'], 
     'correct_idx': 0, 
     'solution': ['Shell-and-tube heat exchanger notation: X-Y where X = shell passes, Y = tube passes', '1-2 configuration: Single shell pass, tubes make 2 passes (U-tube or floating head)', 'Advantages: Compact design, higher heat transfer coefficient (higher velocity in tubes)', 'Disadvantages: Not true counterflow, requires correction factor F for LMTD', 'Common configurations: 1-1, 1-2, 1-4, 1-6, 2-4', 'True counterflow (1-1): Best temperature approach but needs special design (floating head)', 'Selection: Based on required duty, pressure drop, temperature approach, cost. 1-2 is most common industrial configuration. TEMA standards define mechanical design.'], 
     'diagram': False},
    
    # Q85 - Temperature Cross
    {'question': ['Temperature cross in a heat exchanger means:'], 
     'options': ['Hot outlet temperature < Cold outlet temperature', 'Hot outlet temperature > Cold outlet temperature', 'Hot inlet = Cold inlet', 'Occurs in parallel flow only'], 
     'correct_idx': 0, 
     'solution': ['Temperature cross: Hot fluid outlet temperature becomes lower than cold fluid outlet temperature', 'Th,out < Tc,out (This violates second law of thermodynamics if in perfect contact!)', 'Reality: Possible in multi-pass or crossflow exchangers where fluids don\'t have perfect countercurrent contact', 'Cannot occur in true counterflow (1-1 shell-and-tube)', 'LMTD becomes negative or undefined - must use correction factor F or NTU method', 'Example: Hot fluid cools from 200°C to 80°C while cold heats from 50°C to 100°C (cross at 80-100°C)', 'Design implication: Indicates exchanger may not be suitable for that service. Better to redesign for higher approach temperature.'], 
     'diagram': False},
    
    # Q86 - Thermal Expansion
    {'question': ['Thermal expansion joints in pipelines are provided to:'], 
     'options': ['Reduce pressure drop', 'Accommodate length changes due to temperature variation', 'Increase flow rate', 'Prevent corrosion'], 
     'correct_idx': 1, 
     'solution': ['Thermal expansion: Materials expand when heated, contract when cooled. ΔL = α·L·ΔT', 'where α = coefficient of linear expansion, L = original length, ΔT = temperature change', 'For long pipelines with large temperature swings: Significant length change causes stress on equipment and supports', 'Expansion joints: Bellows, slip joints, loops (U-bend, Z-bend) absorb dimensional changes', 'Without expansion joints: Excessive stress → pipe buckling, equipment misalignment, anchor/support failure, leaks', 'Design: Calculate expansion, provide guided supports, use proper anchors and guides', 'Example: 100m carbon steel pipe, ΔT = 100°C, α = 12×10⁻⁶/°C → ΔL = 12cm! Critical in steam lines, furnace piping, long transfer lines.'], 
     'diagram': False},
    
    # Q87 - Liquid-Liquid Extraction
    {'question': ['In liquid-liquid extraction, increasing the number of equilibrium stages (for same solvent rate) generally:'], 
     'options': ['Decreases solute recovery', 'Increases solute recovery', 'Has no effect on recovery', 'Makes extraction impossible'], 
     'correct_idx': 1, 
     'solution': ['Liquid-liquid extraction: Separation based on differential solubility in two immiscible solvents', 'More stages → More contact opportunities → Better mass transfer → Higher recovery of solute', 'Similar to distillation trays: Each stage brings streams closer to equilibrium', 'However: Beyond certain stages, improvement diminishes (approach equilibrium limit)', 'Distribution coefficient K = Concentration in extract / Concentration in raffinate', 'Optimal design: Economic balance between stages (capital cost) and solvent rate (operating cost)', 'Applications: Acetic acid recovery, pharmaceutical purification, nuclear fuel reprocessing. Mixer-settlers or column contactors (RDC, Pulsed, Packed) used.'], 
     'diagram': False},
    
    # Q88 - Leaching
    {'question': ['Leaching is a separation process where:'], 
     'options': ['Gas dissolves in liquid', 'Liquid removes soluble component from solid', 'Two liquids are separated', 'Vapor is condensed'], 
     'correct_idx': 1, 
     'solution': ['Leaching (Solid-liquid extraction): Solvent dissolves soluble component from solid matrix', 'Process: Solid contacted with solvent → Solute dissolves → Separation of liquid (leachate) from solid residue', 'Driving force: Concentration difference between solid and bulk liquid', 'Factors affecting rate: (1) Solvent selection, (2) Temperature, (3) Particle size, (4) Agitation, (5) Contact time', 'Equipment: Batch tanks, continuous extractors (belt, rotary, percolation)', 'Examples: Tea/coffee making, vegetable oil extraction, sugar from sugar beet, metal extraction from ores (copper, gold)', 'Countercurrent operation: Fresh solvent contacts nearly extracted solid → most efficient solvent usage. Important in mineral processing and food industries.'], 
     'diagram': False},
    
    # Q89 - Plug Flow Reactor
    {'question': ['For an isothermal plug flow reactor with first-order kinetics, the relation between conversion X and space time τ is:'], 
     'options': ['X = 1 - e^(-kτ)', 'X = kτ/(1 + kτ)', 'X = e^(-kτ)', 'X = kτ'], 
     'correct_idx': 0, 
     'solution': ['Plug Flow Reactor (PFR): No mixing in flow direction, uniform conditions at any cross-section. All elements have same residence time.', 'For first-order reaction: -rA = kCA', 'Material balance: dCA/dτ = -kCA where τ = space time = V/v₀', 'Integration: CA = CA0·e^(-kτ)', 'Conversion: X = 1 - CA/CA0 = 1 - e^(-kτ)', 'Compare with CSTR: τ = X/[k(1-X)] → PFR is more efficient', 'For high conversion (X → 1): PFR needs finite time but CSTR needs infinite time. PFR used in tubular reactors for gas-phase and high-conversion applications.'], 
     'diagram': False},
    
    # Q90 - Batch Reactor
    {'question': ['For a reversible first-order reaction A ⇌ B in a batch reactor, at long times the concentration of A approaches:'], 
     'options': ['Zero', 'Initial value CA0', 'Equilibrium value determined by equilibrium constant', 'Infinity'], 
     'correct_idx': 2, 
     'solution': ['Reversible reaction: A ⇌ B with forward rate kf and reverse rate kr', 'At equilibrium: Forward rate = Reverse rate, kf·CA,eq = kr·CB,eq', 'Equilibrium constant: K = kf/kr = CB,eq/CA,eq', 'As t → ∞: System approaches equilibrium (not complete conversion)', 'Final CA depends on: (1) Initial conditions, (2) Equilibrium constant K, (3) Stoichiometry', 'If CA0 + CB0 = C (total), then CA,eq = C/(1+K) and CB,eq = KC/(1+K)', 'Unlike irreversible reactions where CA → 0, reversible reactions reach equilibrium state. Important in esterification, isomerization, ammonia synthesis.'], 
     'diagram': False},
    
    # Q91 - Control System Transfer Function
    {'question': ['The standard form of a second-order underdamped control system transfer function is:'], 
     'options': ['ωn²/(s² + 2ζωns + ωn²)', '1/(1 + τs)', 'K/s', 'K/(1 + τs)²'], 
     'correct_idx': 0, 
     'solution': ['Second-order system: G(s) = ωn²/(s² + 2ζωns + ωn²)', 'where ωn = natural frequency, ζ = damping ratio', 'Damping ratio determines response: ζ < 1: Underdamped (oscillatory), ζ = 1: Critically damped, ζ > 1: Overdamped', 'Underdamped (ζ < 1): System oscillates before settling. Overshoot occurs.', 'Characteristics: Rise time, settling time, overshoot depend on ζ and ωn', 'Time constant relation: τ = 1/(ζωn) for dominant pole', 'Example: Spring-mass-damper, level control, temperature control with thermal capacity. Design target: Often ζ = 0.7 (optimal compromise between speed and overshoot).'], 
     'diagram': False},
    
    # Q92 - Ziegler-Nichols Tuning
    {'question': ['In Ziegler-Nichols closed-loop tuning based on ultimate gain Ku and ultimate period Pu, the proportional gain for a classic PID controller is taken as approximately:'], 
     'options': ['0.1Ku', '0.6Ku', 'Ku', '1.2Ku'], 
     'correct_idx': 1, 
     'solution': ['Ziegler-Nichols closed-loop method (Ultimate cycle method):', 'Step 1: Use proportional-only control, increase Kp until sustained oscillations occur', 'Step 2: Note Ku (ultimate gain at oscillation) and Pu (oscillation period)', 'Step 3: Calculate PID parameters: Kp = 0.6Ku, Ti = Pu/2, Td = Pu/8', 'Results in fairly aggressive tuning (often 25% overshoot)', 'Alternative: Cohen-Coon, IMC tuning for less aggressive response', 'ZN gives good starting point but usually requires fine-tuning for specific process. Widely used in industry due to simplicity. Works best for processes with dominant first-order plus dead time behavior.'], 
     'diagram': False},
    
    # Q93 - Inverse Response
    {'question': ['A process shows inverse response when:'], 
     'options': ['Output initially moves in the same direction as final value', 'Output does not change with input', 'Output initially moves opposite to its final steady-state value', 'Process is unstable'], 
     'correct_idx': 2, 
     'solution': ['Inverse response (non-minimum phase): Output initially moves in opposite direction before moving to final value', 'Transfer function has right-half-plane zero (positive zero)', 'Example: Increasing heat to drum boiler → Water expands → Level drops initially (seems counterintuitive!) → Then rises', 'Another example: Increasing reflux in distillation → Column cools → Top product purity drops temporarily → Then improves', 'Challenge for control: Initial wrong-way movement can cause controller to take wrong action', 'Requires: Slow controller tuning (detuning), possibly model predictive control', 'Also called "wrong-way response". Makes process difficult to control manually. Requires understanding of process dynamics for safe operation.'], 
     'diagram': False},
    
    # Q94 - Heat Exchanger Fouling Factor
    {'question': ['For a shell-and-tube heat exchanger, if both shell-side and tube-side fouling factors increase, the overall heat transfer coefficient U will:'], 
     'options': ['Increase', 'Decrease', 'Remain unchanged', 'Become infinite'], 
     'correct_idx': 1, 
     'solution': ['Overall heat transfer coefficient: 1/U = 1/hi + Rf,i + Rwall + Rf,o + 1/ho', 'where Rf,i and Rf,o are fouling resistances (inside and outside)', 'Fouling adds thermal resistance (like adding insulation)', 'If Rf increases → 1/U increases → U decreases', 'Lower U means: (1) Reduced heat transfer rate Q = U·A·ΔT, (2) Worse performance, (3) Need more area or cleaning', 'Typical clean U: 500-1500 W/m²K (water-water). After fouling: May drop to 300-800 W/m²K', 'Design includes fouling allowance: Oversize exchanger by 20-40% to account for fouling over time. Regular cleaning restores performance.'], 
     'diagram': False},
    
    # Q95 - Absorption Column Flooding
    {'question': ['In an absorption column, flooding is characterized by:'], 
     'options': ['Very low gas velocity with no liquid movement', 'Excessively high gas velocity causing liquid accumulation and high pressure drop', 'No pressure drop across packing', 'Complete drying of packing'], 
     'correct_idx': 1, 
     'solution': ['Flooding: Condition where liquid cannot flow down due to high upward gas velocity', 'At flooding: Gas prevents liquid drainage → Liquid accumulates → Pressure drop increases sharply → Poor mass transfer', 'Flooding velocity depends on: Packing type, liquid/gas density ratio, liquid/gas flow rate ratio, packing size', 'Generalized pressure drop correlation (GPDC): Predicts flooding based on flow parameter X = (L/G)√(ρG/ρL)', 'Design: Operate at 60-80% of flooding velocity for stability and turndown', 'Loading point: Point before flooding where pressure drop starts increasing rapidly (70-80% of flooding)', 'Prevention: Proper packing selection, adequate liquid distributor. Flooding limits column capacity in gas treating, CO₂ capture, acid gas removal.'], 
     'diagram': False},
    
    # Q96 - LMTD Correction Factor
    {'question': ['For a log-mean temperature difference (LMTD) calculation in a counter-current heat exchanger, if both terminal temperature differences are equal, the LMTD equals:'], 
     'options': ['Arithmetic mean of the two differences', 'Geometric mean of the two differences', 'The common temperature difference', 'Zero'], 
     'correct_idx': 2, 
     'solution': ['LMTD = (ΔT₁ - ΔT₂)/ln(ΔT₁/ΔT₂) where ΔT₁ and ΔT₂ are terminal temperature differences', 'Special case: If ΔT₁ = ΔT₂ = ΔT', 'LMTD = (ΔT - ΔT)/ln(ΔT/ΔT) = 0/0 (indeterminate form)', 'Apply L\'Hôpital\'s rule or direct limit: LMTD → ΔT', 'Physical meaning: When temperature difference is constant throughout exchanger, LMTD equals that constant difference', 'This occurs in: (1) Condensing vapor at constant T, (2) Boiling liquid at constant T, (3) Perfectly balanced counterflow', 'Simplifies design: Q = U·A·ΔT (no logarithmic mean needed). Common in evaporators and condensers.'], 
     'diagram': False},
    
    # Q97 - Spray Dryer
    {'question': ['In a spray dryer, the predominant mechanism for moisture removal from droplets is:'], 
     'options': ['Filtration', 'Centrifugation', 'Convective evaporation', 'Osmosis'], 
     'correct_idx': 2, 
     'solution': ['Spray drying: Liquid feed atomized into fine droplets, contacted with hot gas, rapid evaporation produces dry powder', 'Mechanism: Convective heat transfer from hot gas to droplet surface → Evaporation of moisture', 'Process: (1) Atomization (nozzle, rotary atomizer), (2) Droplet-gas contact, (3) Evaporation, (4) Product separation (cyclone)', 'Advantages: (1) Fast drying (seconds), (2) Controlled particle size, (3) Suitable for heat-sensitive materials (short exposure)', 'Disadvantages: (1) High energy consumption, (2) Large equipment, (3) Low thermal efficiency (50-70%)', 'Applications: Milk powder, instant coffee, detergents, pharmaceuticals, ceramic powders. Inlet T: 150-250°C, Outlet T: 70-120°C depending on product heat sensitivity.'], 
     'diagram': False},
    
    # Q98 - Packed Tower
    {'question': ['For a packed tower operating with gas-liquid contact, the main reason for using smaller packing size is to:'], 
     'options': ['Increase pressure drop and decrease interfacial area', 'Decrease interfacial area', 'Increase interfacial area and improve mass transfer', 'Increase tower diameter'], 
     'correct_idx': 2, 
     'solution': ['Packing provides surface area for vapor-liquid contact. Surface area per unit volume (a) is key parameter.', 'Smaller packing → Higher surface area per unit volume (a ∝ 1/dp)', 'Higher "a" → More interfacial area → Better mass transfer → Higher efficiency (more theoretical stages per meter)', 'Trade-offs: Smaller packing also means: (1) Higher pressure drop (important in vacuum), (2) Higher cost, (3) More prone to fouling/plugging', 'Common sizes: 25mm (1 inch), 50mm (2 inch), 75mm (3 inch)', 'Selection: Small packing (< 25mm) for high efficiency, large diameter columns. Large packing (> 50mm) for low ΔP, fouling services', 'Modern structured packing: Higher efficiency than random packing. Used in distillation, absorption, stripping columns.'], 
     'diagram': False},
    
    # Q99 - Polyethylene Production
    {'question': ['In polymerization, low-density polyethylene (LDPE) is typically produced by:'], 
     'options': ['Bulk polymerization of vinyl chloride at low pressure', 'High-pressure free-radical polymerization of ethylene', 'Ziegler-Natta catalyzed polymerization of propylene', 'Condensation polymerization of terephthalic acid and ethylene glycol'], 
     'correct_idx': 1, 
     'solution': ['LDPE: Low-Density Polyethylene. Highly branched structure → Lower density (0.91-0.93 g/cm³), crystallinity ~50%', 'Production: Free radical polymerization of ethylene at HIGH pressure (1000-3000 bar) and temperature (80-300°C)', 'Initiator: Peroxide (O₂ or organic peroxide)', 'Properties: Flexible, tough, good moisture barrier. Used in film, bags, squeeze bottles.', 'HDPE: High-Density Polyethylene. Linear chain → Higher density (0.94-0.97 g/cm³), crystallinity ~80%. Made by Ziegler-Natta or Phillips process at LOW pressure.', 'Other polymers mentioned: PVC from vinyl chloride, PP from propylene (Ziegler-Natta), PET from terephthalic acid + ethylene glycol. LDPE was first commercial polyethylene (1930s by ICI).'], 
     'diagram': False},
    
    # Q100 - Payback Period
    {'question': ['For a project with an initial investment of 10 million INR and uniform annual net cash inflow of 2 million INR (no salvage value, ignoring time value of money), the simple payback period is:'], 
     'options': ['2 years', '3 years', '5 years', '10 years'], 
     'correct_idx': 2, 
     'solution': ['Simple payback period: Time required to recover initial investment from net cash inflows', 'Formula: Payback period = Initial Investment / Annual Cash Inflow', 'Given: Initial investment = 10 million INR, Annual inflow = 2 million INR/year', 'Payback = 10/2 = 5 years', 'Assumptions: (1) Uniform cash flows, (2) No time value of money (no discounting)', 'Limitations: (1) Ignores cash flows after payback, (2) Ignores profitability, (3) No consideration of risk', 'Better methods: NPV (Net Present Value), IRR (Internal Rate of Return), Discounted payback. Simple payback useful for quick screening but not sufficient for investment decision. Typical acceptable payback: 2-5 years depending on industry and risk.'], 
     'diagram': False},
]

print("="*70)
print("Chemistry/Engineering - SAIL Questions")
print("Q76-100 with Detailed Solutions")
print("FINAL CENTERED FORMAT (No a,b,c,d prefix in options)")
print("="*70)

doc = Document()
sections = doc.sections
for section in sections:
    section.top_margin = Cm(3.5)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

diagram_questions = []
for i, q_data in enumerate(questions_data):
    q_num = 76 + i
    print(f"Adding Q{q_num}...")
    if q_data.get('diagram', False):
        diagram_questions.append(q_num)
    add_question_table(doc, q_data['question'], q_data['options'], 
                      q_data['correct_idx'], q_data['solution'], 
                      has_diagram=q_data.get('diagram', False))
    if i < len(questions_data) - 1:
        doc.add_page_break()

doc.save('/mnt/user-data/outputs/Chemistry_Engineering_SAIL_Q76-100.docx')

print("\n" + "="*70)
print("✅ DOCUMENT COMPLETED!")
print("="*70)
print(f"✓ Subject: Chemistry/Engineering - SAIL")
print(f"✓ Questions: 25 (Q76-Q100)")
print(f"✓ Solutions: Detailed (4-6 lines minimum)")
print(f"✓ Options: WITHOUT (a), (b), (c), (d) prefix")
print(f"✓ Diagrams: {len(diagram_questions)} questions")
print(f"✓ Table: CENTER ALIGNED (Same format)")
print(f"✓ Page margins: 2.4cm L/R, 3.5cm top")
print(f"✓ Column widths: 1.5cm, 8.5cm, 3.0cm")
if diagram_questions:
    print(f"\n⚠️  Questions with DIAGRAMS: {diagram_questions}")
print("="*70)
print("\n🎉🎉🎉 ALL 100 QUESTIONS COMPLETED! 🎉🎉🎉")