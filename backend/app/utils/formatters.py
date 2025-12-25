"""
Word Document Formatting Utilities.

Ported from existing scripts (b.py, chemical question.py).
Handles cell borders, margins, and text formatting for Word documents.
"""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Union


def set_cell_border(cell, size: str = '4', color: str = '000000'):
    """
    Set cell borders for Word table.

    Args:
        cell: Word table cell object
        size: Border size (4-12, default 4)
        color: Border color in hex (default black: 000000)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_element = OxmlElement(f'w:{edge}')
        edge_element.set(qn('w:val'), 'single')
        edge_element.set(qn('w:sz'), size)
        edge_element.set(qn('w:space'), '0')
        edge_element.set(qn('w:color'), color)
        tcBorders.append(edge_element)

    tcPr.append(tcBorders)


def set_cell_margins(
    cell,
    top: int = 120,
    bottom: int = 120,
    left: int = 100,
    right: int = 100
):
    """
    Set cell margins (padding) for Word table.

    Args:
        cell: Word table cell object
        top: Top margin in twips (default 120, ~6pt)
        bottom: Bottom margin in twips (default 120, ~6pt)
        left: Left margin in twips (default 100, ~5pt)
        right: Right margin in twips (default 100, ~5pt)

    Note: 1 twip = 1/20 of a point, so 20 twips = 1 point
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcMar = OxmlElement('w:tcMar')
    for margin_name, value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)

    tcPr.append(tcMar)


def set_column_width(column, width_cm: float):
    """
    Set column width for table.

    Args:
        column: Table column object
        width_cm: Width in centimeters
    """
    for cell in column.cells:
        cell.width = Cm(width_cm)


def add_text_to_cell(
    cell,
    text_parts: Union[str, List[str]],
    font_name: str = 'Times New Roman',
    font_size: int = 14,
    bold: bool = True,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
):
    """
    Add text to cell with proper formatting.

    Supports multiple paragraphs if text_parts is a list.

    Args:
        cell: Word table cell object
        text_parts: String or list of strings (each becomes a paragraph)
        font_name: Font family (default: Times New Roman)
        font_size: Font size in points (default: 14)
        bold: Bold text (default: True)
        alignment: Text alignment (default: JUSTIFY)
    """
    cell.text = ''

    if isinstance(text_parts, str):
        text_parts = [text_parts]

    for idx, part in enumerate(text_parts):
        if idx > 0:
            paragraph = cell.add_paragraph()
        else:
            paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()

        paragraph.alignment = alignment

        if idx > 0:
            paragraph.paragraph_format.space_before = Pt(3)

        run = paragraph.add_run(part)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold


def format_cell_text(
    cell,
    font_name: str = 'Times New Roman',
    font_size: int = 14,
    bold: bool = True,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
):
    """
    Format existing text in cell.

    Args:
        cell: Word table cell object
        font_name: Font family
        font_size: Font size in points
        bold: Bold text
        alignment: Text alignment
    """
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold


def apply_table_formatting(
    table,
    font_name: str = 'Times New Roman',
    font_size: int = 14,
    bold: bool = True,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
):
    """
    Apply formatting to all cells in table.

    Args:
        table: Word table object
        font_name: Font family
        font_size: Font size in points
        bold: Bold text
        alignment: Text alignment
    """
    for row in table.rows:
        for cell in row.cells:
            format_cell_text(cell, font_name, font_size, bold, alignment)


def add_question_table(
    doc,
    question_parts: Union[str, List[str]],
    options: List[str],
    correct_option_idx: int,
    solution_parts: Union[str, List[str]],
    has_diagram: bool = False
):
    """
    Add a complete question table to document.

    Creates 8-row × 3-column table with:
    - Row 0: Question text
    - Row 1: Type (Multiple_choice)
    - Rows 2-5: Options
    - Row 6: Solution
    - Row 7: Marks

    Args:
        doc: Word Document object
        question_parts: Question text (str or list of paragraphs)
        options: List of 4 option strings
        correct_option_idx: Index of correct option (0-3)
        solution_parts: Solution text (str or list of paragraphs)
        has_diagram: Whether question has diagram

    Returns:
        Created table object
    """
    from docx.enum.table import WD_TABLE_ALIGNMENT

    # Create table: 8 rows × 3 columns
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    set_column_width(table.columns[0], 1.5)   # Label column
    set_column_width(table.columns[1], 8.5)   # Content column
    set_column_width(table.columns[2], 3.0)   # Status column

    # Apply borders and margins to all cells
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell, top=120, bottom=120, left=100, right=100)

    # Row 0: Question
    table.rows[0].cells[0].text = "Question"
    cell = table.rows[0].cells[1].merge(table.rows[0].cells[2])

    if has_diagram:
        q_parts = question_parts if isinstance(question_parts, list) else [question_parts]
        q_parts = q_parts + ['', '[DIAGRAM PRESENT - See PDF]']
        add_text_to_cell(cell, q_parts)
    else:
        add_text_to_cell(cell, question_parts)

    # Row 1: Type
    table.rows[1].cells[0].text = "Type"
    table.rows[1].cells[1].text = "Multiple_choice"
    table.rows[1].cells[1].merge(table.rows[1].cells[2])

    # Rows 2-5: Options
    for i in range(4):
        row_idx = i + 2
        table.rows[row_idx].cells[0].text = "Option"
        table.rows[row_idx].cells[1].text = options[i] if i < len(options) else ""

        if i == correct_option_idx:
            table.rows[row_idx].cells[2].text = "correct"
        else:
            table.rows[row_idx].cells[2].text = "incorrect"

    # Row 6: Solution
    table.rows[6].cells[0].text = "Solution"
    cell = table.rows[6].cells[1].merge(table.rows[6].cells[2])
    add_text_to_cell(cell, solution_parts)

    # Row 7: Marks
    table.rows[7].cells[0].text = "Marks"
    table.rows[7].cells[1].text = "1"
    table.rows[7].cells[2].text = "0.25"

    # Format all cells
    apply_table_formatting(table)

    return table
