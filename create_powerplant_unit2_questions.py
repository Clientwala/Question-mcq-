from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '12')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), '000000')
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)

def set_cell_spacing(cell, spacing=6):
    """Set cell paragraph spacing"""
    for paragraph in cell.paragraphs:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(spacing)
        paragraph_format.space_after = Pt(spacing)

# Create document
doc = Document()

# Set page margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Questions data (Unit 2: Boiler Mountings & Accessories, Q94-108)
questions = [
    {
        "num": 94,
        "question": "A device used to put off fire in the furnace of the boiler when the level of the water in the boiler fails to an unsafe limit.",
        "options": {
            "(a)": "Economiser",
            "(b)": "Fusible plug",
            "(c)": "Superheater",
            "(d)": "Blow off cock"
        },
        "answer": "(b)",
        "solution": "Fusible plug is used to extinguish the fire inside the boiler when the water level inside the boiler falls to an unsafe level and prevent an explosion. It also prevents the damage that may happen due to the explosion."
    },
    {
        "num": 95,
        "question": "Locomotive type of boiler is—",
        "options": {
            "(a)": "horizontal multi-tubular water tube boiler",
            "(b)": "water wall enclosed furnace type",
            "(c)": "vertical tubular fire tube type",
            "(d)": "horizontal multi-tubular fire tube type"
        },
        "answer": "(d)",
        "solution": "Locomotive type of boiler is horizontal multitubular fire tube type. Pressure: 10-20 bar, Capacity: 2-4 tonn/hr. No long chimney is used, artificial draught is achieved by steam jet."
    },
    {
        "num": 96,
        "question": "The Benson boiler has—",
        "options": {
            "(a)": "no steam drum",
            "(b)": "a vertical steam drum",
            "(c)": "an horizontal steam drum",
            "(d)": "two drums-one for water and another for steam"
        },
        "answer": "(a)",
        "solution": "Benson boiler is a modern high pressure, forced circulation, artificial draught, water tube boiler. It was invented by Mark Benson in 1923. This boiler has no steam drum due to very high pressure, upto or more than critical pressure. At critical pressure, latent heat (Δh) = 0, so no water steam separator (like drum) is required."
    },
    {
        "num": 97,
        "question": "Which of the following is a water tube boiler?",
        "options": {
            "(a)": "Cochran boiler",
            "(b)": "Lancashire boiler",
            "(c)": "Babcock Wilcox boiler",
            "(d)": "Locomotive boiler"
        },
        "answer": "(c)",
        "solution": "Water Tube Boiler: In this boiler the water is contained inside the tubes (called water tubes) which are surrounded by flames and hot gases from outside. Examples: Babcock and Wilcox boiler, Stirling boiler, La-Mont boiler, Benson boiler, Yarrow boiler and Loeffler boiler etc."
    },
    {
        "num": 98,
        "question": "In a boiler, the device in which the waste heat of flue gases is utilised for heating feed water is called a/an:",
        "options": {
            "(a)": "superheater",
            "(b)": "injector",
            "(c)": "economiser",
            "(d)": "air preheater"
        },
        "answer": "(c)",
        "solution": "Economiser is a boiler accessory in which the waste heat of flue gases is utilised for heating feed water. It improves boiler efficiency, by 15 to 20% of coal saving. Superheater is used to increase temperature of saturated steam without raising its pressure or at constant pressure."
    },
    {
        "num": 99,
        "question": "The draught in locomotive boilers is produced by",
        "options": {
            "(a)": "Chimney",
            "(b)": "Centrifugal fan",
            "(c)": "Steam jet",
            "(d)": "Locomotion"
        },
        "answer": "(c)",
        "solution": "The draught in locomotive boiler is produced by steam jet. Locomotive boiler is a multi-tubular horizontal, internally fired and mobile boiler. The principle feature of this boiler is to produce steam at a very high rate."
    },
    {
        "num": 100,
        "question": "What can you say about a water tube boiler's efficiency when compared to a fire tube boiler with equal features?",
        "options": {
            "(a)": "Fire tubes has higher efficiency than water tube",
            "(b)": "Water tube has higher efficiency than fire tube",
            "(c)": "Both have the same efficiency",
            "(d)": "Cannot be quantified"
        },
        "answer": "(b)",
        "solution": "In water-tube boiler, water is flowing through tubes which are surrounded by flue gases. These boilers are mostly high pressured. They have higher heat transfer area and higher pressure as compared to fire tube boilers with equal features. So, these water tube boilers have higher efficiency than fire tube boilers."
    },
    {
        "num": 101,
        "question": "The device used to empty the boiler, when required and to discharge the mud, scale of sediments which are accumulated at the bottom of the boiler is known as",
        "options": {
            "(a)": "Safety valve",
            "(b)": "Stop valve",
            "(c)": "Fusible valve",
            "(d)": "Blow off cock"
        },
        "answer": "(d)",
        "solution": "The device used to empty the boiler, when required and to discharge the mud, scale of sediments which are accumulated at the bottom of the boiler is known as blow off cock."
    },
    {
        "num": 102,
        "question": "Which one of the following is a boiler mounting?",
        "options": {
            "(a)": "Fusible plug",
            "(b)": "Super heater",
            "(c)": "Economiser",
            "(d)": "Chimney"
        },
        "answer": "(a)",
        "solution": "Boiler mounting: Boiler mountings are a set of safety device installed for the safe operation of a boiler. Examples: Water level indicator, Safety valve, Pressure gauge, Steam stop valve, Feed check valve, Man hole, Fusible plug. Fusible plug is a boiler mounting."
    },
    {
        "num": 103,
        "question": "Lancashire boiler is:",
        "options": {
            "(a)": "stationary fire tube boiler",
            "(b)": "internally fire tube boiler",
            "(c)": "horizontal boiler",
            "(d)": "All of these"
        },
        "answer": "(d)",
        "solution": "Characteristic of Lancashire boiler: H-FT-NC-ND-IF-2T-LP-S-MP where, H = Horizontal, FT = Fire Tube, NC = Natural Draught, IF = Internal Fired, LP = Low Pressure, S = Stationary, MP = Multipass, 2T = Two Tube. Thermal efficiency = 80 to 90%, Length = 7 to 9 meter, Diameter = 2 to 3 meter."
    },
    {
        "num": 104,
        "question": "A device used to increase the temperature of saturated steam without raising its pressure is called",
        "options": {
            "(a)": "fusible plug",
            "(b)": "blow off cock",
            "(c)": "economiser",
            "(d)": "superheater"
        },
        "answer": "(d)",
        "solution": "A device used to increase the temperature of saturated steam without raising its pressure is called super heater. Superheater takes heat from the flue gases. It is located in the path of flue gases between boiler & economiser. It increases the overall efficiency of cycle. It reduces the moisture content in last stage of turbine and thus increase the work output."
    },
    {
        "num": 105,
        "question": "Which one of the following boiler accessories does not need 'Flue-gas' for its operation?",
        "options": {
            "(a)": "Economiser",
            "(b)": "Pre-heater",
            "(c)": "Injector",
            "(d)": "Super heater"
        },
        "answer": "(c)",
        "solution": "Flue-gas is needed for the operation of following boiler accessories: Economiser, Pre-heater, Super heater. Injector: Flue gas does not need for the operation of injector. The function of an injector is to feed water into the boiler. Its commonly employed for vertical and locomotive boiler and does not find its applications in large capacity high pressure boiler."
    },
    {
        "num": 106,
        "question": "In a water-tube boiler, the flue gas flows through the following accessories: 1. Superheater, 2. ID Fan, 3. Air heater, 4. Economiser. Which of the following gives the correct sequence of the flue gas through these accessories?",
        "options": {
            "(a)": "3, 1, 4, 2",
            "(b)": "1, 4, 3, 2",
            "(c)": "1, 3, 2, 4",
            "(d)": "1, 2, 3, 4"
        },
        "answer": "(b)",
        "solution": "In a water tube boiler the flue gas flows through the following accessories: Superheater → Economiser → Air heater → ID Fan. The correct sequence is 1, 4, 3, 2."
    },
    {
        "num": 107,
        "question": "Which statement represents the steady flow energy equation for a boiler?",
        "options": {
            "(a)": "h₁ + V₁²/2g + Q = h₂ + V₂²/2g",
            "(b)": "Q = h₂ - h₁",
            "(c)": "h₁ + V₁²/2g + Q = h₂ + V₂²/2g",
            "(d)": "W = h₂ - h₁ + Q"
        },
        "answer": "(b)",
        "solution": "For Boiler: W = 0, z₁ = z₂, V₁ = V₂, h₁, h₂ = enthalpies at inlet and outlet. Applying SFEE: h₁ + Q = h₂, Therefore Q = h₂ - h₁"
    },
    {
        "num": 108,
        "question": "What is the primary function of a steam trap in a boiler?",
        "options": {
            "(a)": "Recover heat from exit gases",
            "(b)": "Increase temperature of steam above saturated temperature",
            "(c)": "Separate the water particles from the steam",
            "(d)": "Drain off condensed water accumulating in the steam pipelines"
        },
        "answer": "(d)",
        "solution": "In boiler, primary function of steam trap is to drain off condensed water accumulating in the steam pipelines. A steam trap is a device used to discharge condensates and non-condensable gases with a negligible consumption or loss of live steam."
    }
]

# Create table for each question
for idx, q_data in enumerate(questions):
    # Create an 8-row, 3-column table
    table = doc.add_table(rows=8, cols=3)
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Set column widths
    table.columns[0].width = Cm(2.5)
    table.columns[1].width = Cm(10.5)
    table.columns[2].width = Cm(3.0)

    # Row 1: Question Number, Question Text, "Solution"
    row1 = table.rows[0]
    cell_num = row1.cells[0]
    cell_question = row1.cells[1]
    cell_solution_header = row1.cells[2]

    cell_num.text = f"Q.{q_data['num']}"
    cell_question.text = q_data['question']
    cell_solution_header.text = "Solution"

    # Format Row 1
    for cell in [cell_num, cell_question, cell_solution_header]:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.font.bold = True
        set_cell_spacing(cell, 6)

    # Rows 2-5: Options (a), (b), (c), (d)
    for i, (opt_key, opt_value) in enumerate(q_data['options'].items()):
        row = table.rows[i + 1]
        cell_label = row.cells[0]
        cell_option = row.cells[1]
        cell_status = row.cells[2]

        cell_label.text = opt_key
        cell_option.text = opt_value

        # Mark correct/incorrect
        if opt_key == q_data['answer']:
            cell_status.text = "correct"
        else:
            cell_status.text = "incorrect"

        # Format cells
        for cell in [cell_label, cell_option, cell_status]:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                    run.font.bold = True
            set_cell_spacing(cell, 6)

    # Row 6: Marks row
    row6 = table.rows[5]
    cell_marks_label = row6.cells[0]
    cell_marks_content = row6.cells[1]
    cell_marks_value = row6.cells[2]

    cell_marks_label.text = "Marks:"
    cell_marks_content.text = ""
    cell_marks_value.text = "1 / 0.25"

    for cell in [cell_marks_label, cell_marks_content, cell_marks_value]:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.font.bold = True
        set_cell_spacing(cell, 6)

    # Row 7: Answer row
    row7 = table.rows[6]
    cell_answer_label = row7.cells[0]
    cell_answer_content = row7.cells[1]
    cell_answer_value = row7.cells[2]

    cell_answer_label.text = "Answer:"
    cell_answer_content.text = q_data['answer']
    cell_answer_value.text = ""

    for cell in [cell_answer_label, cell_answer_content, cell_answer_value]:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.font.bold = True
        set_cell_spacing(cell, 6)

    # Row 8: Solution row (merged)
    row8 = table.rows[7]
    cell_sol_label = row8.cells[0]
    cell_sol_content = row8.cells[1]
    cell_sol_empty = row8.cells[2]

    # Merge cells for solution
    cell_merged = cell_sol_label.merge(cell_sol_content).merge(cell_sol_empty)
    cell_merged.text = f"Solution: {q_data['solution']}"

    for paragraph in cell_merged.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.bold = True
    set_cell_spacing(cell_merged, 6)

    # Add borders to all cells
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": 12},
                bottom={"sz": 12},
                left={"sz": 12},
                right={"sz": 12}
            )

    # Add page break after each question (except the last one)
    if idx < len(questions) - 1:
        doc.add_page_break()

# Save document
doc.save('D:\\Question\\PowerPlant_Chapter2_Unit2_Q94-108.docx')
print("Document created successfully!")
print(f"Total questions: {len(questions)}")
print("File saved as: PowerPlant_Chapter2_Unit2_Q94-108.docx")
