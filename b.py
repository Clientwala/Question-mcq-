#!/usr/bin/env python3
"""
Chemistry/Engineering - SAIL Questions
Q51-75 with Detailed Solutions
FINAL CENTERED FORMAT (No a,b,c,d prefix in options)
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_element = OxmlElement(f'w:{edge}')
        edge_element.set(qn('w:val'), 'single')
        edge_element.set(qn('w:sz'), '4')
        edge_element.set(qn('w:space'), '0')
        edge_element.set(qn('w:color'), '000000')
        tcBorders.append(edge_element)
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=120, bottom=120, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_column_width(column, width_cm):
    for cell in column.cells:
        cell.width = Cm(width_cm)

def add_text_to_cell(cell, text_parts):
    cell.text = ''
    if isinstance(text_parts, str):
        text_parts = [text_parts]
    for idx, part in enumerate(text_parts):
        if idx > 0:
            paragraph = cell.add_paragraph()
        else:
            paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if idx > 0:
            paragraph.paragraph_format.space_before = Pt(3)
        run = paragraph.add_run(part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True

def add_question_table(doc, question_parts, options, correct_option_idx, solution_parts, has_diagram=False):
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    set_column_width(table.columns[0], 1.5)
    set_column_width(table.columns[1], 8.5)
    set_column_width(table.columns[2], 3.0)
    
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell, top=120, bottom=120, left=100, right=100)
    
    table.rows[0].cells[0].text = "Question"
    cell = table.rows[0].cells[1].merge(table.rows[0].cells[2])
    if has_diagram:
        add_text_to_cell(cell, question_parts + ['', '[DIAGRAM PRESENT - See PDF]'])
    else:
        add_text_to_cell(cell, question_parts)
    
    table.rows[1].cells[0].text = "Type"
    table.rows[1].cells[1].text = "Multiple_choice"
    table.rows[1].cells[1].merge(table.rows[1].cells[2])
    
    for i in range(4):
        row_idx = i + 2
        table.rows[row_idx].cells[0].text = "Option"
        table.rows[row_idx].cells[1].text = options[i] if i < len(options) else ""
        if i == correct_option_idx:
            table.rows[row_idx].cells[2].text = "correct"
        else:
            table.rows[row_idx].cells[2].text = "incorrect"
    
    table.rows[6].cells[0].text = "Solution"
    cell = table.rows[6].cells[1].merge(table.rows[6].cells[2])
    add_text_to_cell(cell, solution_parts)
    
    table.rows[7].cells[0].text = "Marks"
    table.rows[7].cells[1].text = "1"
    table.rows[7].cells[2].text = "0.25"
    
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                    run.font.bold = True

questions_data = [
    # Q51 - Reaction Order
    {'question': ['For a reaction with rate law: -rA = kCA², the reaction order with respect to A is:'], 
     'options': ['Zero order', 'First order', 'Second order', 'Cannot be determined'], 
     'correct_idx': 2, 
     'solution': ['Reaction order is the exponent of concentration in the rate law.', 'Given rate law: -rA = kCA²', 'The exponent of CA is 2, therefore reaction is second order with respect to A.', 'Overall reaction order = sum of all concentration exponents', 'For -rA = kCA^m CB^n, overall order = m + n', 'Reaction order is determined experimentally and may differ from stoichiometric coefficients. Important for reactor design and kinetic analysis.'], 
     'diagram': False},
    
    # Q52 - Half Life
    {'question': ['For a first-order reaction, the half-life is:'], 
     'options': ['Proportional to initial concentration', 'Inversely proportional to initial concentration', 'Independent of initial concentration', 'Proportional to square of initial concentration'], 
     'correct_idx': 2, 
     'solution': ['Half-life (t1/2) is time required for concentration to reduce to half its initial value.', 'For first-order reaction: -dCA/dt = kCA', 'Integration gives: CA = CA0·e^(-kt)', 'At t = t1/2: CA = CA0/2', 'Solving: t1/2 = ln(2)/k = 0.693/k', 'Notice: t1/2 is independent of CA0 (initial concentration)', 'This is unique property of first-order kinetics. Used in radioactive decay, drug metabolism, and chemical kinetics.'], 
     'diagram': False},
    
    # Q53 - Arrhenius Equation
    {'question': ['The Arrhenius equation relates rate constant to:'], 
     'options': ['Pressure only', 'Temperature only', 'Concentration only', 'Both temperature and activation energy'], 
     'correct_idx': 3, 
     'solution': ['Arrhenius equation: k = A·e^(-Ea/RT)', 'where k = rate constant, A = pre-exponential factor, Ea = activation energy, R = gas constant, T = absolute temperature', 'Shows exponential dependence of k on temperature', 'Higher temperature → higher k → faster reaction', 'Higher Ea → more sensitive to temperature changes', 'Taking ln: ln(k) = ln(A) - Ea/RT (linear form)', 'Plot of ln(k) vs 1/T gives straight line with slope = -Ea/R. Used to determine activation energy and predict reaction rates at different temperatures.'], 
     'diagram': False},
    
    # Q54 - Selectivity
    {'question': ['In parallel reactions A → B (desired) and A → C (undesired), selectivity is defined as:'], 
     'options': ['Rate of B formation / Rate of C formation', 'Concentration of B / Concentration of C', 'Yield of B', 'All of the above'], 
     'correct_idx': 0, 
     'solution': ['Selectivity measures how much desired product forms relative to undesired product.', 'Instantaneous selectivity = rB/rC (ratio of formation rates)', 'Overall selectivity = Amount of B formed / Amount of C formed', 'High selectivity means more desired product, less waste', 'Selectivity can be improved by: (1) Temperature control, (2) Proper catalyst selection, (3) Optimal residence time', 'Different from conversion (how much reactant consumed) and yield (desired product/reactant fed). Critical in fine chemicals and pharmaceuticals to minimize byproducts.'], 
     'diagram': False},
    
    # Q55 - Gas Laws
    {'question': ['For an ideal gas at constant temperature, if pressure is doubled, volume becomes:'], 
     'options': ['Double', 'Half', 'Four times', 'One-fourth'], 
     'correct_idx': 1, 
     'solution': ['For ideal gas: PV = nRT', 'At constant T and n (fixed amount of gas): PV = constant (Boyle\'s Law)', 'Initial state: P1V1 = constant', 'Final state: P2V2 = constant', 'Therefore: P1V1 = P2V2', 'If P2 = 2P1, then V2 = P1V1/P2 = P1V1/(2P1) = V1/2', 'Volume becomes half when pressure doubles. This inverse relationship is fundamental in gas compression and expansion processes.'], 
     'diagram': False},
    
    # Q56 - Psychrometry
    {'question': ['In psychrometry, relative humidity is defined as:'], 
     'options': ['Partial pressure of water / Total pressure', 'Partial pressure of water / Vapor pressure of water at that temperature', 'Humidity ratio / Saturation humidity ratio', 'Temperature / Dew point temperature'], 
     'correct_idx': 1, 
     'solution': ['Relative humidity (RH) = (Partial pressure of water vapor / Saturation vapor pressure at same temperature) × 100%', 'RH = (Pw/Pw,sat) × 100%', 'RH indicates how close air is to saturation', 'RH = 0%: Completely dry air', 'RH = 100%: Saturated air (dew point = dry bulb temperature)', 'Related concepts: Absolute humidity (mass of water/mass of dry air), Dew point (temperature at which condensation begins)', 'Important in HVAC design, drying operations, and comfort analysis. Psychrometric chart displays these relationships graphically.'], 
     'diagram': False},
    
    # Q57 - Crystallization
    {'question': ['In crystallization, supersaturation is the driving force for:'], 
     'options': ['Crystal dissolution', 'Both nucleation and crystal growth', 'Temperature increase', 'Solvent evaporation'], 
     'correct_idx': 1, 
     'solution': ['Supersaturation = (Actual concentration - Equilibrium solubility)/Equilibrium solubility', 'Supersaturation creates thermodynamic driving force for crystallization.', 'Two phenomena occur: (1) Nucleation: Formation of new crystal nuclei, (2) Crystal growth: Existing crystals grow larger', 'Higher supersaturation → more nucleation → smaller crystals', 'Lower supersaturation → more growth on existing crystals → larger crystals', 'Control of supersaturation is key to controlling crystal size distribution.', 'Achieved by cooling, evaporation, or antisolvent addition. Critical in pharmaceutical manufacturing and chemical purification.'], 
     'diagram': False},
    
    # Q58 - Filtration
    {'question': ['In filtration, the pressure drop across the filter is proportional to:'], 
     'options': ['Filtration rate', 'Square of filtration rate', 'Inverse of filtration rate', 'Cube of filtration rate'], 
     'correct_idx': 0, 
     'solution': ['Filtration follows Darcy\'s law (for incompressible cakes): -ΔP = (μαwc/A²)V + (μRm/A)V̇', 'where ΔP = pressure drop, V̇ = filtration rate (dV/dt), μ = viscosity, α = specific cake resistance', 'For constant pressure filtration: Linear relationship between ΔP and average V̇', 'For compressible cakes: α increases with ΔP, making relationship more complex', 'Pressure drop increases with: (1) Higher flow rate, (2) Thicker cake, (3) Higher viscosity, (4) Smaller particles', 'Types: Plate-and-frame, rotary vacuum, cartridge filters. Used in pharmaceutical, food, and wastewater treatment industries.'], 
     'diagram': False},
    
    # Q59 - Fluidization
    {'question': ['Minimum fluidization velocity is the velocity at which:'], 
     'options': ['Particles start to settle', 'Pressure drop equals weight of bed', 'Particles are completely entrained', 'Bubbling begins'], 
     'correct_idx': 1, 
     'solution': ['Fluidization occurs when upward drag force on particles equals their weight.', 'At minimum fluidization velocity (Umf): Pressure drop = Weight of bed/Area', 'Below Umf: Fixed bed (packed bed flow)', 'At Umf: Incipient fluidization (bed just starts to expand)', 'Above Umf: Fluidized bed (particles suspended, liquid-like behavior)', 'Ergun equation predicts Umf: Accounts for particle size, density, void fraction, fluid properties', 'Applications: Fluidized bed reactors (FCC, polymerization), dryers, combustors. Advantages: Excellent mixing, heat transfer, mass transfer.'], 
     'diagram': False},
    
    # Q60 - Size Reduction
    {'question': ['Rittinger\'s law for crushing and grinding states that energy required is proportional to:'], 
     'options': ['New surface area created', 'Volume of material crushed', 'Mass of material crushed', 'Cube of particle diameter'], 
     'correct_idx': 0, 
     'solution': ['Rittinger\'s law: Energy required is proportional to new surface area created.', 'E = KR(1/L2 - 1/L1) where L1 = feed size, L2 = product size, KR = Rittinger\'s constant', 'Applicable for fine grinding where surface area changes are significant.', 'Other laws: Kick\'s law (energy ∝ reduction ratio), Bond\'s law (intermediate, most widely used)', 'Size reduction increases surface area for: (1) Chemical reactions, (2) Heat/mass transfer, (3) Mixing', 'Equipment: Jaw crushers, ball mills, hammer mills, fluid energy mills. Important in mining, cement, pharmaceuticals, and food processing.'], 
     'diagram': False},
    
    # Q61 - Centrifugal Separation
    {'question': ['In a centrifuge, separation is enhanced by increasing:'], 
     'options': ['Gravitational acceleration', 'Centrifugal acceleration (RPM)', 'Temperature', 'Pressure'], 
     'correct_idx': 1, 
     'solution': ['Centrifugal force = mω²r where m = mass, ω = angular velocity, r = radius', 'Centrifugal acceleration ac = ω²r = (2πN)²r where N = rotational speed (RPS)', 'Separation factor = ac/g = ω²r/g (how many times stronger than gravity)', 'Higher RPM → higher ac → faster and better separation', 'Stokes\' law settling velocity increases proportionally with acceleration', 'Applications: (1) Removing suspended solids from liquids, (2) Separating immiscible liquids, (3) Cream separation from milk', 'Types: Tubular, disc, decanter centrifuges. Used in pharmaceutical, food, chemical, and wastewater industries.'], 
     'diagram': False},
    
    # Q62 - pH Definition
    {'question': ['pH is defined as:'], 
     'options': ['log[H⁺]', '-log[H⁺]', 'ln[H⁺]', '-ln[H⁺]'], 
     'correct_idx': 1, 
     'solution': ['pH = -log₁₀[H⁺] where [H⁺] = hydrogen ion concentration in mol/L', 'pH scale: 0 to 14 (in water at 25°C)', 'pH < 7: Acidic solution (higher [H⁺])', 'pH = 7: Neutral solution ([H⁺] = [OH⁻] = 10⁻⁷ M)', 'pH > 7: Basic/alkaline solution (lower [H⁺])', 'Relation to pOH: pH + pOH = 14 (at 25°C), pOH = -log[OH⁻]',