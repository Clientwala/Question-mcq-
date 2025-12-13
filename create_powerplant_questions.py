from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '12')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), '000000')
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)

# Create document
doc = Document()

# Set page margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(3.5)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

# Questions data with solutions from PDF
questions = [
    {
        "question": "Thermal power plant works on",
        "options": [
            {"text": "Carnot cycle", "correct": False},
            {"text": "Joule cycle", "correct": False},
            {"text": "Rankine cycle", "correct": True},
            {"text": "Otto cycle", "correct": False}
        ],
        "solution": """Thermal power plant works on Rankine cycle.

Cycle Applications:
• Carnot cycle - Used for comparison of other cycles
• Rankine cycle - Thermal power plant
• Joule cycle - Gas Turbine
• Otto cycle - Petrol engine
• Diesel cycle - Low speed diesel engine
• Bell Coleman cycle - Air Refrigerator"""
    },
    {
        "question": "Rankine cycle has:",
        "options": [
            {"text": "two isochoric and two isobaric processes", "correct": False},
            {"text": "two isothermal and two isobaric processes", "correct": False},
            {"text": "two isentropic and two isobaric processes", "correct": True},
            {"text": "two adiabatic and two isobaric processes", "correct": False}
        ],
        "solution": """Rankine cycle has two isentropic and two isobaric processes.

Process breakdown:
• 5-1-2: Heat addition at constant pressure in boiler
• 2-3: Isentropic work done by turbine
• 3-4: Heat rejection at constant pressure in condenser
• 4-5: Isentropic pump work

Steam power plant is based on Rankine cycle."""
    },
    {
        "question": "In cooling tower, water is cooled by the process of:",
        "options": [
            {"text": "Condensation", "correct": False},
            {"text": "Fusion", "correct": False},
            {"text": "Evaporation", "correct": True},
            {"text": "Sublimation", "correct": False}
        ],
        "solution": """A cooling tower is a special heat exchanger in which air and water are brought into direct contact with each other in order to reduce the water's temperature.

The purpose of cooling tower is to cool relatively warm water by contacting with unsaturated air. The evaporation of water mainly provides cooling."""
    },
    {
        "question": "Bleeding in turbine means-",
        "options": [
            {"text": "Leakage of steam", "correct": False},
            {"text": "Steam doing no useful work", "correct": False},
            {"text": "Extracting steam for pre-heating feed water", "correct": True},
            {"text": "Removal of condenser steam", "correct": False}
        ],
        "solution": """Bleeding is the process of extracting live steam from certain point and utilizing heat of this steam in raising the temperature of water entering into boiler.

For this purpose feed water heater (FWH) is used."""
    },
    {
        "question": "Expansion of steam in Rankine cycle is assumed to be",
        "options": [
            {"text": "isentropic", "correct": True},
            {"text": "polytropic", "correct": False},
            {"text": "isothermal", "correct": False},
            {"text": "hyperbolic", "correct": False}
        ],
        "solution": """Expansion of steam in Rankine cycle is assumed to be isentropic process.

Rankine Cycle is an ideal cycle for steam power plant. It has four processes - Two isobaric & two reversible adiabatic or isentropic processes.

Heat addition & rejection in Rankine cycle is at constant pressure & compression & expansion of steam perform at constant entropy (isentropically).

Processes:
(2-3) = Isentropic expansion (Turbine)
(3-4) = Isobaric heat rejection (Condenser)
(4-5) = Isentropic compression (Pump)
(5-1-2) = Isobaric heat addition (Boiler)"""
    },
    {
        "question": "Between Rankine and Carnot cycles, which of the following is not a difference?",
        "options": [
            {"text": "Carnot cycle is theoretical whereas Rankine cycle is practical", "correct": False},
            {"text": "Carnot cycle uses air whereas Rankine cycle uses water as a working substance", "correct": False},
            {"text": "Carnot cycle exchanges heat at constant temperature whereas Rankine cycle exchanges heat at constant pressure", "correct": False},
            {"text": "Carnot cycle is ideal for vapor power cycles whereas Rankine cycle is ideal for heat engines", "correct": True}
        ],
        "solution": """Carnot cycle and Rankine cycle both are theoretical cycles, but Carnot cycle has highest efficiency compared to other cycles. So this cycle is considered as ideal cycle and used for the comparison of efficiency of other cycles.

Key differences:
• As compared to Carnot cycle, Rankine cycle can be considered as practical
• Carnot cycle uses air, whereas Rankine cycle uses water as a working substance, so in Rankine cycle, we find a vapour dome
• Carnot cycle exchanges heat at constant temperature (T = C), whereas Rankine cycle exchanges heat at constant pressure
• Carnot cycle is ideal cycle for all other cycles (Gas power, heat engines, gas negative, vapour power cycle all) whereas Rankine cycle is ideal for vapour power cycles only"""
    },
    {
        "question": "Which process efficiency is stated in terms of the thermal efficiency of a steam turbine determined using the Rankine cycle model?",
        "options": [
            {"text": "Isenthalpic", "correct": False},
            {"text": "Isentropic", "correct": True},
            {"text": "Hydraulic", "correct": False},
            {"text": "Volumetric", "correct": False}
        ],
        "solution": """Using the Rankine cycle, the thermal efficiency of the steam turbines is determined in terms of isentropic process."""
    },
    {
        "question": "In steam power cycle, the process of removing non-condensable gases is called",
        "options": [
            {"text": "scavenging process", "correct": False},
            {"text": "deaeration process", "correct": True},
            {"text": "exhaust process", "correct": False},
            {"text": "condensation process", "correct": False}
        ],
        "solution": """The process of removing non-condensable gases in steam power cycle is called deaeration and is done in deaerator.

Other processes:
• Scavenging process is the process of removing exhaust gases and replacing with fresh gases
• Exhaust process is the removal of burnt out gases
• Condensation process is conversion of steam to liquid form"""
    },
    {
        "question": "A power plant, which uses a gas turbine following by steam turbine for power generation, is called",
        "options": [
            {"text": "Topping cycle", "correct": False},
            {"text": "Bottoming cycle", "correct": False},
            {"text": "Brayton cycle", "correct": False},
            {"text": "Combined cycle", "correct": True}
        ],
        "solution": """A power plant, which uses a gas turbine following by steam turbine for power generation, is called combined cycle."""
    },
    {
        "question": "A regenerative steam cycle renders",
        "options": [
            {"text": "increased work output per unit mass of steam", "correct": False},
            {"text": "decrease work output per unit mass of steam", "correct": False},
            {"text": "increased thermal efficiency", "correct": False},
            {"text": "decreased work output per unit mass of steam as well as increased thermal efficiency", "correct": True}
        ],
        "solution": """A regenerative steam cycle renders decreased work output per unit mass of steam as well as increased thermal efficiency.

A steam turbine cycle in which the condensate or feed water is heated to a temperature that is much higher than that corresponding to saturation at the exhaust pressure by means of steam that has been bled from the turbine at points intermediate between the throttle and exhaust.

Efficiency of Rankine cycle:
η = Wnet/Q_a = (W_T - W_P)/Q_a = (h_2 - h_3) - (h_5 - h_4)/(h_2 - h_5)"""
    },
    {
        "question": "Regenerative heating ........... the thermal efficiency of a Rankine cycle.",
        "options": [
            {"text": "does not affect", "correct": False},
            {"text": "decreases", "correct": False},
            {"text": "may increase or decrease", "correct": False},
            {"text": "increases", "correct": True}
        ],
        "solution": """Regenerative heating increases the thermal efficiency of Rankine cycle.

In regenerative cycle, extract live steam (Bleeding) from certain point and utilizing heat of this steam in raising the temperature of water entering into boiler. For this purpose feed water heater (FWH) is used.

In regenerative cycle, Q_a decreases, So η increases.

Note: Rankine efficiency would approach Carnot cycle efficiency by providing a series of regenerative feed heaters."""
    },
    {
        "question": "The reheat cycle in steam power plant is mainly adopted to:",
        "options": [
            {"text": "increase moisture content in low pressure stages to a safe value", "correct": False},
            {"text": "decrease moisture content in low pressure stage to a safe value", "correct": True},
            {"text": "decrease the capacity of condenser", "correct": False},
            {"text": "recover the waste heat boiler", "correct": False}
        ],
        "solution": """Reheat cycle in steam power plant is mainly adopted to decrease moisture content in low pressure stage to a safe value.

Process breakdown:
3'-4: Isentropic expansion of steam (In high pressure turbine)
4-5: Reheating of steam at constant pressure (In reheater)
5-6: Isentropic expansion of steam (In low pressure turbine)
6-1: Condensation at constant pressure (In condenser)
2-3-3': Heat addition to water at constant pressure (In boiler)

The reheat cycle in steam power plant is mainly adopted to decrease moisture content in low pressure stage to a safe value."""
    },
    {
        "question": "The work output from the turbine in case of a Rankine cycle is given by:",
        "options": [
            {"text": "enthalpy change between inlet and outlet", "correct": True},
            {"text": "entropy change between inlet and outlet", "correct": False},
            {"text": "pressure change between inlet and outlet", "correct": False},
            {"text": "temperature change between inlet and outlet", "correct": False}
        ],
        "solution": """The work output from the turbine in case of a Rankine cycle is given by enthalpy change between inlet and outlet.

Rankine cycle on (h-s) diagram:
W_turbine = (h_2 - h_3)
W_pump = (h_5 - h_4)
Q_a = (h_2 - h_5)

Net work output = W_T - W_P = (h_2 - h_3) - (h_5 - h_4)

Work ratio, r_w = W_net/W_T = (W_T - W_P)/W_T"""
    },
    {
        "question": "For a steam power plant ______ cycle is recommended.",
        "options": [
            {"text": "Brayton", "correct": False},
            {"text": "Carnot", "correct": False},
            {"text": "Rankine", "correct": True},
            {"text": "Otto", "correct": False}
        ],
        "solution": """Rankine cycle is recommended for a steam power plant. Because it is most efficient cycle & practically possible cycle for power plant.

This cycle has two isobaric (heat addition & rejection) & two isentropic (compression & expansion) Process.

η = W_net/Q_a = (T_H - T_C)/T_H where W_net = W_T - W_P and Q_a = Q_b + Q_r

Note:
• Brayton cycle - For gas turbine or gas power plant
• Otto cycle - For S.I. engine"""
    },
    {
        "question": "An increase in the efficiency of Rankine cycle can be expected with",
        "options": [
            {"text": "increase in exhaust pressure", "correct": False},
            {"text": "increase in temperature of heat rejection", "correct": False},
            {"text": "decrease in exhaust pressure", "correct": True},
            {"text": "decrease in temperature of heat addition", "correct": False}
        ],
        "solution": """An increase in the efficiency of Rankine cycle can be expected with decrease in exhaust pressure.

Method of improving Rankine cycle efficiency (η_R):
1. By increasing boiler pressure (P_b)
2. By increasing superheat temperature (t_sp)
3. By decreasing condenser or exhaust pressure (P_c)

The efficiency increases when:
• Increasing the boiler pressure
• Superheating the steam to high temperature
• Decreasing the exhaust/condenser pressure"""
    }
]

# Create table for each question
for idx, q_data in enumerate(questions, 1):
    # Create a table with 8 rows and 3 columns
    table = doc.add_table(rows=8, cols=3)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Set column widths
    table.columns[0].width = Cm(2.5)
    table.columns[1].width = Cm(10.5)
    table.columns[2].width = Cm(3.0)

    # Row 0: Question row
    row0 = table.rows[0]
    cell_q_label = row0.cells[0]
    cell_q_text = row0.cells[1]
    cell_q_right = row0.cells[2]

    # Merge columns 2 and 3 for question row
    cell_q_text.merge(cell_q_right)

    cell_q_label.text = "Question"
    cell_q_text.text = q_data["question"]

    for cell in [cell_q_label, cell_q_text]:
        para = cell.paragraphs[0]
        run = para.runs[0] if para.runs else para.add_run(cell.text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        para.space_after = Pt(6)

    # Row 1: Type row
    row1 = table.rows[1]
    cell_type_label = row1.cells[0]
    cell_type_text = row1.cells[1]
    cell_type_right = row1.cells[2]

    # Merge columns 2 and 3 for type row
    cell_type_text.merge(cell_type_right)

    cell_type_label.text = "Type"
    cell_type_text.text = "Multiple_choice"

    for cell in [cell_type_label, cell_type_text]:
        para = cell.paragraphs[0]
        run = para.runs[0] if para.runs else para.add_run(cell.text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        para.space_after = Pt(6)

    # Rows 2-5: Options (4 options)
    for opt_idx, option in enumerate(q_data["options"]):
        row = table.rows[2 + opt_idx]
        cell_opt_label = row.cells[0]
        cell_opt_text = row.cells[1]
        cell_opt_status = row.cells[2]

        cell_opt_label.text = "Option"
        cell_opt_text.text = option["text"]
        cell_opt_status.text = "correct" if option["correct"] else "incorrect"

        for cell in [cell_opt_label, cell_opt_text, cell_opt_status]:
            para = cell.paragraphs[0]
            run = para.runs[0] if para.runs else para.add_run(cell.text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.bold = True
            para.space_after = Pt(6)

    # Row 6: Solution row
    row6 = table.rows[6]
    cell_sol_label = row6.cells[0]
    cell_sol_text = row6.cells[1]
    cell_sol_right = row6.cells[2]

    # Merge columns 2 and 3 for solution row
    cell_sol_text.merge(cell_sol_right)

    cell_sol_label.text = "Solution"
    cell_sol_text.text = q_data["solution"]

    for cell in [cell_sol_label, cell_sol_text]:
        para = cell.paragraphs[0]
        run = para.runs[0] if para.runs else para.add_run(cell.text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        para.space_after = Pt(6)

    # Row 7: Marks row
    row7 = table.rows[7]
    cell_marks_label = row7.cells[0]
    cell_marks_val = row7.cells[1]
    cell_marks_points = row7.cells[2]

    cell_marks_label.text = "Marks"
    cell_marks_val.text = "1"
    cell_marks_points.text = "0.25"

    for cell in [cell_marks_label, cell_marks_val, cell_marks_points]:
        para = cell.paragraphs[0]
        run = para.runs[0] if para.runs else para.add_run(cell.text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        para.space_after = Pt(6)

    # Add borders to all cells
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": 12},
                bottom={"sz": 12},
                left={"sz": 12},
                right={"sz": 12}
            )

    # Add page break after each question
    doc.add_page_break()

# Save document
doc.save('D:\\Question\\PowerPlant_Chapter2_Unit1_Q1-15_Final.docx')
print("Document created successfully!")
print("File saved as: PowerPlant_Chapter2_Unit1_Q1-15_Final.docx")
