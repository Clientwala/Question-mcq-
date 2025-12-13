from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    Set cell borders
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Create borders element
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '12')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), '000000')
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)

# Create document
doc = Document()

# Set page margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(3.5)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

# Questions data
questions = [
    {
        "num": 1,
        "question": "Thermal power plant works on",
        "options": {
            "(a)": "Carnot cycle",
            "(b)": "Joule cycle",
            "(c)": "Rankine cycle",
            "(d)": "Otto cycle"
        },
        "answer": "(c)"
    },
    {
        "num": 2,
        "question": "Rankine cycle has:",
        "options": {
            "(a)": "two isochoric and two isobaric processes",
            "(b)": "two isothermal and two isobaric processes",
            "(c)": "two isentropic and two isobaric processes",
            "(d)": "two adiabatic and two isobaric processes"
        },
        "answer": "(c)"
    },
    {
        "num": 3,
        "question": "In cooling tower, water is cooled by the process of:",
        "options": {
            "(a)": "Condensation",
            "(b)": "Fusion",
            "(c)": "Evaporation",
            "(d)": "Sublimation"
        },
        "answer": "(c)"
    },
    {
        "num": 4,
        "question": "Bleeding in turbine means-",
        "options": {
            "(a)": "Leakage of steam",
            "(b)": "Steam doing no useful work",
            "(c)": "Extracting steam for pre-heating feed water",
            "(d)": "Removal of condenser steam"
        },
        "answer": "(c)"
    },
    {
        "num": 5,
        "question": "Expansion of steam in Rankine cycle is assumed to be",
        "options": {
            "(a)": "isentropic",
            "(b)": "polytropic",
            "(c)": "isothermal",
            "(d)": "hyperbolic"
        },
        "answer": "(a)"
    },
    {
        "num": 6,
        "question": "Between Rankine and Carnot cycles, which of the following is not a difference?",
        "options": {
            "(a)": "Carnot cycle is theoretical whereas Rankine cycle is practical",
            "(b)": "Carnot cycle uses air whereas Rankine cycle uses water as a working substance",
            "(c)": "Carnot cycle exchanges heat at constant temperature whereas Rankine cycle exchanges heat at constant pressure",
            "(d)": "Carnot cycle is ideal for vapor power cycles whereas Rankine cycle is ideal for heat engines"
        },
        "answer": "(d)"
    },
    {
        "num": 7,
        "question": "Which process efficiency is stated in terms of the thermal efficiency of a steam turbine determined using the Rankine cycle model?",
        "options": {
            "(a)": "Isenthalpic",
            "(b)": "Isentropic",
            "(c)": "Hydraulic",
            "(d)": "Volumetric"
        },
        "answer": "(b)"
    },
    {
        "num": 8,
        "question": "In steam power cycle, the process of removing non-condensable gases is called",
        "options": {
            "(a)": "scavenging process",
            "(b)": "deaeration process",
            "(c)": "exhaust process",
            "(d)": "condensation process"
        },
        "answer": "(b)"
    },
    {
        "num": 9,
        "question": "A power plant, which uses a gas turbine following by steam turbine for power generation, is called",
        "options": {
            "(a)": "Topping cycle",
            "(b)": "Bottoming cycle",
            "(c)": "Brayton cycle",
            "(d)": "Combined cycle"
        },
        "answer": "(d)"
    },
    {
        "num": 10,
        "question": "A regenerative steam cycle renders",
        "options": {
            "(a)": "increased work output per unit mass of steam",
            "(b)": "decrease work output per unit mass of steam",
            "(c)": "increased thermal efficiency",
            "(d)": "decreased work output per unit mass of steam as well as increased thermal efficiency"
        },
        "answer": "(d)"
    },
    {
        "num": 11,
        "question": "Regenerative heating ........... the thermal efficiency of a Rankine cycle.",
        "options": {
            "(a)": "does not affect",
            "(b)": "decreases",
            "(c)": "may increase of decrease",
            "(d)": "increases"
        },
        "answer": "(d)"
    },
    {
        "num": 12,
        "question": "The reheat cycle in steam power plant is mainly adopted to:",
        "options": {
            "(a)": "increase moisture content in low pressure stages to a safe value",
            "(b)": "decrease moisture content in low pressure stage to a safe value",
            "(c)": "decrease the capacity of condenser",
            "(d)": "recover the waste heat boiler"
        },
        "answer": "(b)"
    },
    {
        "num": 13,
        "question": "The work output from the turbine in case of a Rankine cycle is given by:",
        "options": {
            "(a)": "enthalpy change between inlet and outlet",
            "(b)": "entropy change between inlet and outlet",
            "(c)": "pressure change between inlet and outlet",
            "(d)": "temperature change between inlet and outlet"
        },
        "answer": "(a)"
    },
    {
        "num": 14,
        "question": "For a steam power plant ______ cycle is recommended.",
        "options": {
            "(a)": "Brayton",
            "(b)": "Carnot",
            "(c)": "Rankine",
            "(d)": "Otto"
        },
        "answer": "(c)"
    },
    {
        "num": 15,
        "question": "An increase in the efficiency of Rankine cycle can be expected with",
        "options": {
            "(a)": "increase in exhaust pressure",
            "(b)": "increase in temperature of heat rejection",
            "(c)": "decrease in exhaust pressure",
            "(d)": "decrease in temperature of heat addition"
        },
        "answer": "(c)"
    }
]

# Create table for each question
for q_data in questions:
    # Create a table with 3 rows and 3 columns
    table = doc.add_table(rows=3, cols=3)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Set column widths
    table.columns[0].width = Cm(1.5)
    table.columns[1].width = Cm(8.5)
    table.columns[2].width = Cm(3.0)

    # Row 1: Question
    row1 = table.rows[0]
    cell_num = row1.cells[0]
    cell_question = row1.cells[1]
    cell_solution_header = row1.cells[2]

    # Question number
    cell_num.text = str(q_data["num"])
    para = cell_num.paragraphs[0]
    para.runs[0].font.name = 'Times New Roman'
    para.runs[0].font.size = Pt(14)
    para.runs[0].font.bold = True

    # Question text
    cell_question.text = q_data["question"]
    para = cell_question.paragraphs[0]
    para.runs[0].font.name = 'Times New Roman'
    para.runs[0].font.size = Pt(14)
    para.runs[0].font.bold = True

    # Solution header
    cell_solution_header.text = "Solution"
    para = cell_solution_header.paragraphs[0]
    para.runs[0].font.name = 'Times New Roman'
    para.runs[0].font.size = Pt(14)
    para.runs[0].font.bold = True

    # Row 2: Options
    row2 = table.rows[1]
    cell_empty = row2.cells[0]
    cell_options = row2.cells[1]
    cell_answer = row2.cells[2]

    # Empty cell
    cell_empty.text = ""

    # Options
    options_text = ""
    for key, value in q_data["options"].items():
        marker = "✓ " if key == q_data["answer"] else ""
        options_text += f"{marker}{key} {value}\n"
    cell_options.text = options_text.strip()
    para = cell_options.paragraphs[0]
    para.runs[0].font.name = 'Times New Roman'
    para.runs[0].font.size = Pt(14)
    para.runs[0].font.bold = True

    # Answer
    cell_answer.text = q_data["answer"]
    para = cell_answer.paragraphs[0]
    para.runs[0].font.name = 'Times New Roman'
    para.runs[0].font.size = Pt(14)
    para.runs[0].font.bold = True

    # Row 3: Solution row (merged)
    row3 = table.rows[2]
    cell_sol_label = row3.cells[0]
    cell_sol_text = row3.cells[1]
    cell_sol_ans = row3.cells[2]

    # Merge cells for solution
    cell_sol_label.merge(cell_sol_text).merge(cell_sol_ans)
    cell_sol_label.text = f"Answer: {q_data['answer']}"
    para = cell_sol_label.paragraphs[0]
    para.runs[0].font.name = 'Times New Roman'
    para.runs[0].font.size = Pt(14)
    para.runs[0].font.bold = True

    # Add borders to all cells
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": 12},
                bottom={"sz": 12},
                left={"sz": 12},
                right={"sz": 12}
            )

    # Add space after table
    doc.add_paragraph()

# Save document
doc.save('D:\\Question\\PowerPlant_Chapter2_Unit1_Q1-15.docx')
print("Document created successfully!")
